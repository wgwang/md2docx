import os
import io
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from markdown_it import MarkdownIt
from mdit_py_plugins.texmath import texmath_plugin
import latex2mathml.converter
import mathml2omml
from lxml import etree

class MarkdownToDocx:
    def __init__(self, template_path=None, font_name="Microsoft YaHei", preserve_breaks=False):
        self.font_name = font_name
        self.preserve_breaks = preserve_breaks
        # Enable tables and math
        self.md = MarkdownIt('commonmark').enable('table').use(texmath_plugin)
        
        if template_path:
            self.doc = Document(template_path)
        else:
            self.doc = Document()
        
        if self.font_name:
            style = self.doc.styles['Normal']
            style.font.name = self.font_name
            style.font.size = Pt(11)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)

    def _apply_font(self, run):
        if not self.font_name:
            return
        run.font.name = self.font_name
        # This is critical for CJK fonts to show up correctly in Word
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)

    def convert(self, md_content, output_path):
        # Reset doc for a fresh convert if reused, but typically we instantiate new.
        # For safety, we assume one instance = one convert or we'd need to reset self.doc
        tokens = self.md.parse(md_content)
        self._process_tokens(tokens)
        self.doc.save(output_path)

    def _process_tokens(self, tokens):
        # We'll use a stack for styles/context
        # contexts: 'root', 'list_bullet', 'list_ordered', 'blockquote'
        context_stack = ['root']
        
        i = 0
        while i < len(tokens):
            token = tokens[i]
            
            if token.type == 'heading_open':
                level = int(token.tag[1])
                i += 1
                if i < len(tokens) and tokens[i].type == 'inline':
                    p = self.doc.add_heading(level=level)
                    self._process_inline(p, tokens[i].children)
            
            elif token.type == 'paragraph_open':
                # Determine style based on context
                style = None
                if context_stack[-1] == 'list_bullet':
                    style = 'List Bullet'
                elif context_stack[-1] == 'list_ordered':
                    style = 'List Number'
                elif context_stack[-1] == 'blockquote':
                    style = 'Quote' # or similar
                
                i += 1
                if i < len(tokens) and tokens[i].type == 'inline':
                    if style:
                        # python-docx requires the style to exist in the template usually.
                        # List Bullet and List Number are standard.
                        try:
                            p = self.doc.add_paragraph(style=style)
                        except:
                            p = self.doc.add_paragraph()
                    else:
                        p = self.doc.add_paragraph()
                    self._process_inline(p, tokens[i].children)

            elif token.type == 'bullet_list_open':
                context_stack.append('list_bullet')
            elif token.type == 'bullet_list_close':
                if context_stack[-1] == 'list_bullet': context_stack.pop()

            elif token.type == 'ordered_list_open':
                context_stack.append('list_ordered')
            elif token.type == 'ordered_list_close':
                if context_stack[-1] == 'list_ordered': context_stack.pop()
                
            elif token.type == 'blockquote_open':
                context_stack.append('blockquote')
            elif token.type == 'blockquote_close':
                if context_stack[-1] == 'blockquote': context_stack.pop()

            elif token.type == 'fence' or token.type == 'code_block':
                p = self.doc.add_paragraph()
                # Basic code styling
                run = p.add_run(token.content)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            elif token.type == 'table_open':
                i = self._process_table(tokens, i)
            
            elif token.type == 'math_block':
                p = self.doc.add_paragraph()
                # Center block equations
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                self._add_math(p, token.content)

            elif token.type == 'hr':
                self.doc.add_page_break() # or add a line
                
            i += 1

    def _process_table(self, tokens, start_index):
        # Scan to find table end and dimensions
        table_end_index = start_index
        depth = 0
        rows = 0
        cols = 0
        
        # Determine scope and dimensions
        current_cols = 0
        in_tr = False
        
        for j in range(start_index, len(tokens)):
            t = tokens[j]
            if t.type == 'table_open':
                depth += 1
            elif t.type == 'table_close':
                depth -= 1
                if depth == 0:
                    table_end_index = j
                    break
            
            if depth == 1: # Top level of this table
                if t.type == 'tr_open':
                    rows += 1
                    current_cols = 0
                    in_tr = True
                elif t.type == 'tr_close':
                    in_tr = False
                    if current_cols > cols:
                        cols = current_cols
                elif (t.type == 'th_open' or t.type == 'td_open') and in_tr:
                    current_cols += 1
        
        # Create table if valid
        if rows > 0 and cols > 0:
            try:
                table = self.doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
            except Exception:
                # Fallback if 'Table Grid' style is missing
                table = self.doc.add_table(rows=rows, cols=cols)
            
            row_idx = -1
            col_idx = 0
            
            j = start_index
            while j < table_end_index:
                t = tokens[j]
                if t.type == 'tr_open':
                    row_idx += 1
                    col_idx = 0
                elif t.type == 'th_open' or t.type == 'td_open':
                    # Look ahead for inline content
                    if j + 1 < len(tokens) and tokens[j+1].type == 'inline':
                        if row_idx < rows and col_idx < cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Use the first paragraph of the cell
                            p = cell.paragraphs[0]
                            self._process_inline(p, tokens[j+1].children)
                    col_idx += 1
                j += 1
        
        return table_end_index

    def _process_inline(self, paragraph, children):
        if not children:
            return
            
        # Inline state
        is_bold = False
        is_italic = False
        is_code = False
        
        for child in children:
            if child.type == 'text':
                run = paragraph.add_run(child.content)
                if is_bold: run.bold = True
                if is_italic: run.italic = True
                if is_code: 
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(220, 50, 50) # Red-ish for code
                else:
                    self._apply_font(run)

            elif child.type == 'strong_open':
                is_bold = True
            elif child.type == 'strong_close':
                is_bold = False
            
            elif child.type == 'em_open':
                is_italic = True
            elif child.type == 'em_close':
                is_italic = False
                
            elif child.type == 'code_inline':
                # code_inline token has content directly usually? 
                # Checking markdown-it structure: type='code_inline', content='foo'
                run = paragraph.add_run(child.content)
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(200, 50, 50)
                # Code inline usually doesn't need CJK font fix unless it has CJK, 
                # but good to apply fallback if needed. Courier New might support it or not.
            
            elif child.type == 'math_inline' or child.type == 'math_single':
                self._add_math(paragraph, child.content)
            
            elif child.type == 'image':
                # Image token has src in attrs
                src = child.attrs.get('src')
                if src:
                    self._add_image(paragraph, src)
            
            elif child.type == 'softbreak':
                if self.preserve_breaks:
                    paragraph.add_run().add_break()
                else:
                    paragraph.add_run(' ')
            
            elif child.type == 'hardbreak':
                paragraph.add_run().add_break()
            
            elif child.type == 'link_open':
                pass # docx links are complex (hyperlinks), maybe skip or just text
            elif child.type == 'link_close':
                pass

    def _add_math(self, paragraph, latex_content):
        # Strip potential surrounding delimiters that might have been included
        latex_content = latex_content.strip()
        if latex_content.startswith('$') and latex_content.endswith('$'):
            latex_content = latex_content[1:-1]

        try:
            mathml = latex2mathml.converter.convert(latex_content)
            omml_string = mathml2omml.convert(mathml)
            # Add namespace if missing
            if 'xmlns:m=' not in omml_string:
                omml_string = omml_string.replace('<m:oMath>', '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">', 1)
            
            # Parse and append
            tree = etree.fromstring(omml_string)
            self._clean_omml(tree)
            paragraph._element.append(tree)
        except Exception as e:
            print(f"Failed to convert math: {e}")
            paragraph.add_run(latex_content)

    def _clean_omml(self, tree):
        ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        m_ns = ns['m']
        
        def is_empty(node):
            if node is None:
                return True
            # Check text content recursively
            text = "".join(node.itertext()).strip()
            return not text

        def get_or_create_pr(element, pr_name):
            pr = element.find(f'./m:{pr_name}', ns)
            if pr is None:
                pr = etree.Element(f'{{{m_ns}}}{pr_name}')
                element.insert(0, pr)
            return pr

        def set_hide(pr, hide_name):
            hide = pr.find(f'./m:{hide_name}', ns)
            if hide is None:
                hide = etree.SubElement(pr, f'{{{m_ns}}}{hide_name}')
                hide.set(f'{{{m_ns}}}val', 'on')

        # 1. Fix m:rad with empty m:deg
        for rad in tree.xpath('.//m:rad', namespaces=ns):
            deg = rad.find('./m:deg', ns)
            should_hide_deg = False
            
            if deg is not None:
                if is_empty(deg):
                    rad.remove(deg)
                    should_hide_deg = True
            else:
                should_hide_deg = True
            
            if should_hide_deg:
                radPr = get_or_create_pr(rad, 'radPr')
                set_hide(radPr, 'degHide')

        # 2. Fix m:nary with empty m:sup or m:sub
        for nary in tree.xpath('.//m:nary', namespaces=ns):
            sup = nary.find('./m:sup', ns)
            if sup is None or is_empty(sup):
                if sup is not None:
                    nary.remove(sup)
                naryPr = get_or_create_pr(nary, 'naryPr')
                set_hide(naryPr, 'supHide')

            sub = nary.find('./m:sub', ns)
            if sub is None or is_empty(sub):
                if sub is not None:
                    nary.remove(sub)
                naryPr = get_or_create_pr(nary, 'naryPr')
                set_hide(naryPr, 'subHide')

        # 3. Fix m:sSubSup with empty m:sup -> m:sSub
        for ssubsup in tree.xpath('.//m:sSubSup', namespaces=ns):
            sup = ssubsup.find('./m:sup', ns)
            sub = ssubsup.find('./m:sub', ns)
            
            sup_empty = sup is not None and is_empty(sup)
            sub_empty = sub is not None and is_empty(sub)
            
            if sup_empty:
                ssubsup.remove(sup)
                ssubsup.tag = f'{{{m_ns}}}sSub'
                # If sub is also empty, it will fall through to be a weird sSub with empty sub, which is fine-ish,
                # or we could handle double empty. But prioritized the main issue.
                
            elif sub_empty:
                ssubsup.remove(sub)
                ssubsup.tag = f'{{{m_ns}}}sSup'

    def _add_image(self, paragraph, src):
        # We need to fetch the image
        try:
            image_stream = None
            if src.startswith('http://') or src.startswith('https://'):
                response = requests.get(src, timeout=10)
                response.raise_for_status()
                image_stream = io.BytesIO(response.content)
            else:
                if os.path.exists(src):
                    image_stream = open(src, 'rb')
            
            if image_stream:
                # Add a run with the picture
                run = paragraph.add_run()
                
                # Add picture without specifying width to get natural size
                shape = run.add_picture(image_stream)

                # Calculate available width based on page margins
                try:
                    section = self.doc.sections[0]
                    available_width = section.page_width - section.left_margin - section.right_margin
                    
                    if shape.width > available_width:
                        # Resize to fit available width while maintaining aspect ratio
                        aspect_ratio = shape.height / shape.width
                        shape.width = available_width
                        shape.height = int(available_width * aspect_ratio)
                except Exception as e:
                    print(f"Warning: Could not resize image {src}: {e}")

                if not src.startswith('http'):
                    image_stream.close()
        except Exception as e:
            print(f"Failed to load image {src}: {e}")
            paragraph.add_run(f"[Image: {src}]")
