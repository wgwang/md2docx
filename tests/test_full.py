import unittest
import os
import io
from md2docx import MarkdownToDocx
from docx import Document

class TestMd2DocxFull(unittest.TestCase):
    def setUp(self):
        self.output_file = 'test_full_output.docx'

    def tearDown(self):
        if os.path.exists(self.output_file):
            os.remove(self.output_file)

    def test_basic_elements(self):
        md = """# Heading 1
## Heading 2
Normal text with **bold** and *italic* and `inline code`.

- Bullet 1
- Bullet 2

1. Ordered 1
2. Ordered 2

> Blockquote
"""
        converter = MarkdownToDocx()
        converter.convert(md, self.output_file)
        self.assertTrue(os.path.exists(self.output_file))
        doc = Document(self.output_file)
        # Check headings
        self.assertEqual(doc.paragraphs[0].text, "Heading 1")
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        self.assertEqual(doc.paragraphs[1].text, "Heading 2")
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 2")

    def test_tables(self):
        md = """
| Header 1 | Header 2 |
| -------- | -------- |
| Cell 1,1 | Cell 1,2 |
| Cell 2,1 | Cell 2,2 |
"""
        converter = MarkdownToDocx()
        converter.convert(md, self.output_file)
        doc = Document(self.output_file)
        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3) # Header + 2 data rows
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, "Cell 1,1")

    def test_math(self):
        md = """
Inline math: $E=mc^2$

Block math:
$$
a^2 + b^2 = c^2
$$
"""
        converter = MarkdownToDocx()
        converter.convert(md, self.output_file)
        self.assertTrue(os.path.exists(self.output_file))
        # Hard to verify OXML content without deep dive, but at least it shouldn't crash.

    def test_convert_to_bytes(self):
        md = "# Hello Bytes"
        converter = MarkdownToDocx()
        stream = converter.convert_to_bytes(md)
        self.assertIsInstance(stream, io.BytesIO)
        self.assertGreater(stream.getbuffer().nbytes, 0)
        
        # Verify it's a valid docx
        doc = Document(stream)
        self.assertEqual(doc.paragraphs[0].text, "Hello Bytes")

    def test_templates(self):
        template_path = "professional_template.docx"
        if not os.path.exists(template_path):
             # Generate it if missing for test
             from scripts.generate_template import create_professional_template
             create_professional_template(template_path)
        
        md = "# Professional Doc"
        converter = MarkdownToDocx(template_path=template_path)
        converter.convert(md, self.output_file)
        doc = Document(self.output_file)
        self.assertEqual(doc.paragraphs[0].text, "Professional Doc")
        # Check if header was preserved from template
        self.assertIn("PROFESSIONAL REPORT", doc.sections[0].header.paragraphs[0].text)

if __name__ == '__main__':
    unittest.main()
