import unittest
from md2docx.converter import MarkdownToDocx
import os
from docx import Document

class TestMathBug(unittest.TestCase):
    def setUp(self):
        self.converter = MarkdownToDocx()
        self.output_file = "test_math_bug.docx"

    def tearDown(self):
        if os.path.exists(self.output_file):
            os.remove(self.output_file)

    def test_complex_math_bug(self):
        # The formula reported by the user
        md_content = r"$$ \lambda_{e_s,e_o}^{r} (t | \bar{t}) = f (g_{e_s, e_o}^r (\bar{t}) ) \cdot ( t - \bar{t} ) $$"
        
        # This should not raise an exception now
        try:
            self.converter.convert(md_content, self.output_file)
        except Exception as e:
            self.fail(f"Conversion failed with error: {e}")
        
        self.assertTrue(os.path.exists(self.output_file))
        
        # Verify the content (at least that it can be opened and has some math)
        doc = Document(self.output_file)
        # We can't easily check for OMML structure without more complex parsing,
        # but the fact that it didn't crash and produced a docx is a good sign.
        # The fallback in _add_math is to add the latex string as text if conversion fails,
        # so we should check if the LaTeX string is present in the document
        # ONLY if the conversion failed.
        
        # If conversion succeeded, the text should NOT contain the raw LaTeX (except if we want to be very thorough).
        # Actually, in _add_math:
        # except Exception as e:
        #     print(f"Failed to convert math: {e}")
        #     paragraph.add_run(latex_content)
        
        # So we can check if it's rendered as math (element with 'm' namespace)
        has_math = False
        for p in doc.paragraphs:
            if 'oMath' in p._element.xml:
                has_math = True
                break
        self.assertTrue(has_math, "Document should contain math element")

if __name__ == '__main__':
    unittest.main()
