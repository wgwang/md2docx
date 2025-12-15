import os
import unittest
from docx import Document
from md2docx.converter import MarkdownToDocx

class TestTemplate(unittest.TestCase):
    def setUp(self):
        self.test_md = "test_template.md"
        self.test_docx = "test_output.docx"
        self.template_docx = "test_template.docx"
        
        with open(self.test_md, "w", encoding="utf-8") as f:
            f.write("# Hello World\nThis is a test.")

        # Create a dummy template
        doc = Document()
        doc.save(self.template_docx)

    def tearDown(self):
        if os.path.exists(self.test_md):
            os.remove(self.test_md)
        if os.path.exists(self.test_docx):
            os.remove(self.test_docx)
        if os.path.exists(self.template_docx):
            os.remove(self.template_docx)

    def test_conversion_with_template(self):
        # Run conversion with template
        converter = MarkdownToDocx(template_path=self.template_docx, font_name=None)
        
        with open(self.test_md, "r", encoding="utf-8") as f:
            content = f.read()
        
        converter.convert(content, self.test_docx)
        
        self.assertTrue(os.path.exists(self.test_docx))
        
        # Open the result and check
        doc = Document(self.test_docx)
        self.assertTrue(len(doc.paragraphs) > 0)
        # Check that font was NOT forced to Microsoft YaHei (since we passed None)
        # The first paragraph is Heading 1. Second is 'This is a test.'
        # Let's check the run of the second paragraph.
        p = doc.paragraphs[1]
        if p.runs:
            # If font_name is None, the run.font.name might be None (inheriting)
            self.assertIsNone(p.runs[0].font.name)

    def test_conversion_default_behavior(self):
        # Run conversion without template (should use default font)
        # Note: In library usage, we must pass font_name explicit if we want it, 
        # but here we test the class logic which doesn't default inside __init__ anymore 
        # if we pass explicit values.
        # Wait, I changed __init__ default to "Microsoft YaHei".
        # So MarkdownToDocx() should default to YaHei.
        
        converter = MarkdownToDocx()
        with open(self.test_md, "r", encoding="utf-8") as f:
            content = f.read()
        converter.convert(content, self.test_docx)
        
        doc = Document(self.test_docx)
        # Check that font WAS forced
        p = doc.paragraphs[1]
        if p.runs:
            self.assertEqual(p.runs[0].font.name, "Microsoft YaHei")

if __name__ == "__main__":
    unittest.main()
